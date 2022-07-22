import time
import datetime
import os
import docx
from docx.dml.color import ColorFormat
from docx.shared import RGBColor
from docx.shared import Inches
from docx.text.run import Font, Run
from PIL import ImageGrab
import cyrtranslit
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as ec
from selenium.webdriver.support.ui import WebDriverWait


class MyException(Exception):
    pass


# Вход в приложение


def add_driver():
    s = Service(r"C:\Users\a.zvezdin\Documents\PYTHON\chromedriver.exe")
    driver = webdriver.Chrome(service=s)
    return driver


def login(driver, url, name, passw):
    driver.get(url)
    driver.maximize_window()
    user_name = driver.find_element(By.ID, "loginEdit-el")
    user_name.clear()
    user_name.send_keys(name)
    password = driver.find_element(By.ID, "passwordEdit-el")
    password.clear()
    password.send_keys(passw)
    driver.find_element(By.ID, 't-comp14-textEl').click()
    el = driver.find_elements(By.XPATH, "//div[@class = 'ts-messagebox-caption']")
    dl = len(el)
    if dl > 0:
        text = el[0].text
        print("Появилось окошко с сообщением:\n" + text)
        print('Введен неправильный логин/пароль')
        write_file('Ошибка авторизации', "Error")
        write_file(text, "Error")
        screen_to_file(driver, 100)
        raise MyException('Введен неправильный логин/пароль')
    text = 'Вход под пользователем \"' + name + '\" выполнен'
    write_file(text, "Good")
    #print('Вход под пользователем \"' + name + '\" выполнен')


# Выход из приложения
def exit_user(driver):
    driver.find_element(By.ID, "profile-user-button-wrapperEl").click()
    wait = WebDriverWait(driver, 10)
    el = wait.until(
        ec.element_to_be_clickable((By.XPATH, ".//li[contains(concat(' ', @data-item-marker, ' '), 'Выход')]")))
    el.click()
    # el = wait.until(ec.element_to_be_clickable((By.XPATH,".//span[contains(concat(' ', @data-item-marker, ' '), 'Да')]")))
    el = driver.find_elements(By.XPATH, ".//span[contains(concat(' ', @data-item-marker, ' '), 'Да')]")
    dl = len(el)
    if dl > 0:
        el[0].click()



# Проверка состояния

def action_condition(driver):
    status = driver.find_elements(By.XPATH, "//div[contains(concat(' ', @class, ' '), 'stage-current') and @data-item-marker]")
    if bool(len(status)):
        s = status[0].get_attribute('data-item-marker')
        text = 'Cостояние: \"' + s + '\"'
        write_file(text, "Good")
        print(text)
        return s
    else:
        text = 'Ошибка загрузки страницы'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)



def select_stage(driver, stage):
    status = driver.find_elements(By.XPATH, "//div[contains(concat(' ', @class, ' '), 'stage-item') and @data-item-marker]")
    if bool(len(status)):
        st = driver.find_element(By.XPATH, "//div[@class = 'actions-dashboard-header']").find_elements(By.XPATH, "//div[@data-item-marker = '{0}']".format(stage))
        dl = len(st)
        if dl > 0:
            s = st[0].get_attribute('class')
            dl = s.find('stage-current')
            if dl > 0:
                text = 'Объект уже в состояни \"' + stage + '\"'
                write_file(text, "Error")
                print(text)
                screen_to_file(driver, 100)
                raise MyException(text)
            dl = s.find('stage-with-menu')
            if dl >= 0:
                st[0].click()
                el1 = driver.find_elements(By.XPATH, ".//li[@data-item-marker = '{0}']".format(stage))
                if len(el1) > 0:
                    cl = el1[0].get_attribute('class')
                    dl1 = cl.find('menu-item-disabled')
                    if dl1 >= 0:
                        text = 'Состояние \"' + stage + '\" недоступно для перехода'
                        write_file(text, "Error")
                        print(text)
                        screen_to_file(driver, 100)
                        raise MyException(text)
                    driver.find_element(By.XPATH, ".//li[@data-item-marker = '{0}']".format(stage)).click()
                    text = 'Выбран переход в состояние \"' + stage + '\"'
                    write_file(text, "Good")
                    print(text)
            else:
                dl = s.find('stage-enabled')
                if dl > 0:
                    st[0].click()
                    text = 'Выбран переход на стадию \"' + stage + '\"'
                    write_file(text, "Good")
                    print(text)
                else:
                    text = 'Состояние \"' + stage + '\" недоступно для перехода'
                    write_file(text, "Error")
                    print(text)
                    screen_to_file(driver, 100)
                    raise MyException(text)

        else:
            menu = driver.find_elements(By.XPATH, "//div[contains(concat(' ', @class, ' '), 'stage-with-menu') and @data-item-marker]")
            dl = len(menu)
            if dl > 0:
                print(dl)
                print(menu[0].get_attribute('outerHTML'))
                for num in range(dl):
                    print("01")
                    menu[num].click()
                    print("01")
                    el1 = driver.find_elements(By.XPATH, ".//li[@data-item-marker = '{0}']".format(stage))
                    if len(el1) > 0:
                        cl = el1[0].get_attribute('class')
                        dl1 = cl.find('menu-item-disabled')
                        if dl1 >= 0:
                            text = 'Состояние \"' + stage + '\" недоступно для перехода'
                            write_file(text, "Error")
                            print(text)
                            screen_to_file(driver, 100)
                            raise MyException(text)

                        driver.find_element(By.XPATH, ".//li[@data-item-marker = '{0}']".format(stage)).click()
                        text = 'Выбран переход в состояние \"' + stage + '\"'
                        write_file(text, "Good")
                        print(text)
                        break
                    text = 'Состояние \"' + stage + '\" не найдено'
                    write_file(text, "Error")
                    print(text)
                    screen_to_file(driver, 100)
                    raise MyException(text)
            else:
                text = 'Состояние \"' + stage + '\" не найдено'
                write_file(text, "Error")
                print(text)
                screen_to_file(driver, 100)
                raise MyException(text)

    else:
        text = 'Ошибка загрузки страницы'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)


# Открываем рабочее место
def open_work_place(driver, place):

    el = driver.find_elements(By.XPATH, "//span[@data-item-marker='Развернуть панель разделов']")
    dl = len(el)
    if dl > 0:
        el[0].click()
    WebDriverWait(driver, 10).until(ec.element_to_be_clickable((By.ID, "top-menu-workplace-button-container"))).click()
    el = driver.find_elements(By.XPATH, "//li[@data-item-marker='{0}']".format(place))
    if len(el) == 0:
        text = 'Рабочее место \"' + place + '\" не найдено'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    WebDriverWait(driver, 10).until(
        ec.element_to_be_clickable((By.XPATH, "//li[@data-item-marker='{0}']".format(place)))).click()
    text = 'Выбрано рабочее место \"' + place + '\"'
    write_file(text, "Good")
    print(text)
    time.sleep(2)


# Открываем раздел
def open_section(driver, section):

    el = driver.find_elements(By.XPATH, "//div[@data-item-marker = '{0}' and @class = 'ts-sidebar-item-image']".format(section))
    if len(el) == 0:
        text = 'Раздел \"' + section + '\" не найден'
        write_file(text, "Error")
        print('Раздел \"' + section + '\" не найден')
        screen_to_file(driver, 100)
        raise MyException(text)
    driver.find_element(By.XPATH, "//div[@data-item-marker = '{0}' and @class = 'ts-sidebar-item-image']".format(section)).click()
    text = 'Открыт раздел: \"' + section + '\"'
    write_file(text, "Good")
    print('Открыт раздел: \"' + section + '\"')
    time.sleep(2)


# Прокручивание страницы элементов

def scroll_web_page(driver, id_element):

    dl = len(driver.find_elements(By.XPATH, ".//div[contains(concat(' ', @id, ' '), '{0}')]".format(id_element)))
    while bool(dl):
        driver.find_element(By.TAG_NAME, 'body').send_keys(Keys.END)
        time.sleep(0.5)
        x = driver.find_elements(By.XPATH, ".//div[contains(concat(' ', @id, ' '), '{0}')]".format(id_element))
        if (len(x) - dl) == 0:
            dl = 0
        else:
            dl = len(x)


# Очистка предфильтров
def cleaning_prefilters(driver):

    driver.find_element(By.XPATH,
        ".//span[contains(concat(' ', @id, ' '), 'SectionFilterDropDownMenuButton-menuWrapEl')]").click()
    # driver.implicitly_wait(5)
    driver.find_element(By.XPATH, ".//li[contains(concat(' ', @data-item-marker, ' '), 'фильтры')]").click()


# Очистка созданных фильтров
def cleaning_filters(driver):

    el = driver.find_elements(By.XPATH, ".//span[contains(concat(' ', @data-tag, ' '), 'customFilter')]")
    dl = len(el)
    if bool(dl):
        for key in el:
            key.click()

# Выбор записи в разделе
def select_object(driver, name, action):

    el = driver.find_elements(By.XPATH, ".//div[contains(concat(' ', @data-item-marker, ' '), '{0}')]".format(name))
    dl = len(el)
    if dl == 0:
        text = 'Запись \"' + name + '\" не найдена'
        write_file(text, "Error")
        screen_to_file(driver, 100)
        raise MyException(text)
    if dl > 1:
        text = 'Не удается однозначно определить запись \"' + name + '\", будет выбрана первая.'
        write_file(text, "Error")
        screen_to_file(driver, 100)
        el[0].click()
        el3 = driver.find_element(By.XPATH, ".//div[contains(concat(' ', @data-item-marker, ' '), '{0}')]".format(name))
        el4 = el3.find_elements(By.XPATH, ".//span[@data-item-marker =  '{0}']".format(action))
        dl2 = len(el)
        if dl2 == 0:
            text = 'Для записи \"' + name + '\" действие \"' + action + '\" недоступно'
            write_file(text, "Error")
            screen_to_file(driver, 100)
            raise MyException(text)
        el4[0].click()
        text = 'Для записи \"' + name + '\" выбрано действие \"' + action + '\"'
        write_file(text, "Good")
        screen_to_file(driver, 100)
        print(text)
    else:
        el[0].click()
        el = driver.find_element(By.XPATH, ".//div[contains(concat(' ', @data-item-marker, ' '), '{0}')]".format(name))
        el1 = el.find_elements(By.XPATH, ".//span[@data-item-marker =  '{0}']".format(action))
        dl = len(el1)
        if dl == 0:
            text = 'Для записи \"' + name + '\" действие \"' + action + '\" недоступно'
            write_file(text, "Error")
            screen_to_file(driver, 100)
            raise MyException(text)
        el1[0].click()
        text = 'Для записи \"' + name + '\" выбрано действие \"' + action + '\"'
        write_file(text, "Good")
        print(text)

def click_checkbox(driver, name, action):

    check = driver.find_elements(By.XPATH, ".//div[contains(concat(' ', @class, 'checkbox-label-wrap'), ' ') and label[text() = '{0}']]".format(name))
    dl = len(check)
    if dl == 0:
        text = 'Логическое поле \"' + name + '\" не найдено'
        write_file(text, "Error")
        screen_to_file(driver, 100)
        raise MyException(text)
    if dl > 1:
        text = 'Найдено несколько похожих логических полей \"' + name + '\"'
        write_file(text, "Error")
        screen_to_file(driver, 100)
        raise MyException(text)

    el1 = check[0].find_element(By.XPATH, "..").find_element(By.XPATH, ".//input[@type='checkbox']")
    class_check = el1.find_element(By.XPATH, "..").get_attribute('class')
    dl = class_check.find('t-checkboxedit-checked')
    if action == "Off":
        if dl >= 0:
            check.click()
            text = 'Логическое поле \"' + name + '\" деактивировано'
            write_file(text, "Good")
            print(text)
    if action == "On":
        if dl < 0:
            check[0].click()
            text = 'Логическое поле \"' + name + '\" активировано'
            write_file(text, "Good")
            print(text)


# создание фильтра по наименованию
def create_filters(driver, name, filters):

    #driver.find_element(By.XPATH, ".//span[text() = 'Фильтры/группы']").click()
    click_button(driver, 'Фильтры/группы')
    action_tools_button(driver, "Добавить условие")
    driver.find_element(By.XPATH, ".//input[@placeholder = 'Колонка']").clear()
    driver.find_element(By.XPATH, ".//input[@placeholder = 'Колонка']").send_keys("{0}".format(name))
    action_tools_button(driver, name)
    driver.find_element(By.XPATH, ".//input[@placeholder = 'Значение']").send_keys("{0}".format(filters))
    driver.find_element(By.XPATH, ".//span[@data-item-marker = 'applyButton']").click()
    text = 'Установлен фильтр по колонке: \"' + name +'\" Со значением \"' + filters + '\"'
    write_file(text, "Good")
    print(text)
    time.sleep(1)


# Ввод текстового поля
def enter_value_input(driver, name, value):

    el = driver.find_elements(By.XPATH, "//label[text() = '{0}' and contains(concat(' ', @class, ' '), 't-label')]".format(name))
    dl = len(el)
    if dl == 0:
        text = 'Поле \"' + name + '\" не найдено'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    if dl > 1:
        text = 'Найдено несколько подходящих полей \"' + name + '\"'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    else:
        el1 = el[0].find_element(By.XPATH, "../../..").find_element(By.TAG_NAME, "input")
        if el1.get_attribute('readonly'):
            text = 'Поле \"' + name + '\" недоступно для редактирования'
            write_file(text, "Error")
            print(text)
            screen_to_file(driver, 100)
            raise MyException(text)
        el1.clear()
        el1.send_keys("{0}".format(value))
        text = 'В поле \"' + name + '\" внесено значение: \"' + value + '\"'
        write_file(text, "Good")
        print(text)

# Ввод текстового поля без наименования


def enter_value_noname_input(driver, name, value):

    el = driver.find_elements(By.XPATH, ".//div[contains(concat(' ', @data-item-marker, ' '), '{0}')]".format(name))
    dl = len(el)
    if dl == 0:
        text = 'Поле \"' + name + '\" не найдено'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    if dl > 1:
        text = 'Найдено несколько подходящих полей \"' + name + '\"'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    else:
        el1 = el[0].find_element(By.TAG_NAME, "input")
        if el1.get_attribute('readonly'):
            text = 'Поле \"' + name + '\" недоступно для редактирования'
            write_file(text, "Error")
            print(text)
            screen_to_file(driver, 100)
            raise MyException(text)
        el1.clear()
        el1.send_keys("{0}".format(value))
        text = 'В поле \"' + name + '\" внесено значение: \"' + value + '\"'
        write_file(text, "Good")
        print(text)








# Чтение даныных из текстового поля и поля-справочника

def read_value(driver, name):
    el = driver.find_elements(By.XPATH, "//label[text() = '{0}' and contains(concat(' ', @class, ' '), 't-label t-label-is-required')]".format(name))
    dl = len(el)
    if dl == 0:
        text = 'Поле \"' + name + '\" не найдено'
        write_file(text, "Error")
        screen_to_file(driver, 100)
        raise MyException(text)
    if dl > 1:
        text = 'Найдено несколько подходящих полей \"' + name + '\"'
        print(text)
        write_file(text, "Error")
    else:
        el1 = el[0].find_element(By.XPATH, "../../..")
        value = el1.find_element(By.TAG_NAME, "input").get_attribute('value')
        text = 'В поле \"' + name + '\" содержится значение: \"' + value + '\"'
        write_file(text, "Good")
        print(text)
        return value



# Ввод многострочного текстового поля

def enter_value_textarea(driver,name, value):
    el = driver.find_elements(By.XPATH, "//label[text() = '{0}' and contains(concat(' ', @class, ' '), 't-label')]".format(name))
    dl = len(el)
    if dl == 0:
        text = 'Поле \"' + name + '\" не найдено'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    if dl > 1:
        text = 'Найдено несколько подходящих полей \"' + name + '\"'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    else:
        el1 = el[0].find_element(By.XPATH, "../../..").find_element(By.XPATH, ".//textarea[2]")
        if el1.get_attribute('readonly'):
            text = 'Поле \"' + name + '\" недоступно для редактирования'
            write_file(text, "Error")
            print(text)
            screen_to_file(driver, 100)
            raise MyException(text)
        el1.clear()
        el1.send_keys("{0}".format(value))
        text = 'В поле \"' + name + '\" внесено значение: \"' + value + '\"'
        write_file(text, "Good")
        print(text)

# Чтение многострочного текстового поля
def read_textarea_value(driver, name):
    el = driver.find_elements(By.XPATH,
        "//label[text() = '{0}' and contains(concat(' ', @class, ' '), 't-label')]".format(name))
    dl = len(el)
    if dl == 0:
        text = 'Поле \"' + name + '\" не найдено'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    if dl > 1:
        text = 'Найдено несколько подходящих полей \"' + name + '\"'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    else:
        el1 = el[0].find_element(By.XPATH, "../../..")
        value = el1.find_element(By.TAG_XPATH, ".//textarea[2]").get_attribute('value')
        text = 'В поле \"' + name + '\" содержится значение: \"' + value + '\"'
        write_file(text, "Good")
        print(text)
        return value



# Ввод поля-справочника
def enter_list_value(driver, name, value):

    el = driver.find_elements(By.XPATH, "//label[text() = '{0}' and contains(concat(' ', @class, ' '), 't-label ')]".format(name))
    dl = len(el)
    if dl == 0:
        text = 'Поле \"' + name + '\" не найдено'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    #if dl > 1:
    #    text = 'Не удалось однозначно определить поле \"' + name + '\"'
    #    write_file(text, "Error")
    #    print(text)
    #    screen_to_file(driver, 100)
    #   raise MyException(text)
    else:
        el1 = el[0].find_element(By.XPATH, "../../..").find_element(By.TAG_NAME, "input")
        if el1.get_attribute('readonly'):
            text = 'Поле \"' + name + '\" недоступно для редактирования'
            write_file(text, "Error")
            print(text)
            screen_to_file(driver, 100)
            raise MyException(text)
        clear = el[0].find_element(By.XPATH, "../../..").find_elements(By.TAG_NAME, "a")
        dl = len(clear)
        if dl > 0:
            text = clear[0].get_attribute('href')
            dl = len(text)
            if dl > 0:
                hover = ActionChains(driver).move_to_element(clear[0])
                hover.perform()
                time.sleep(0.5)
                el1 = el[0].find_element(By.XPATH, "../../..").find_element(By.XPATH, ".//div[@class = 'base-edit-clear-icon']")
                el1.click()

        else:
            el1 = el[0].find_element(By.XPATH, "../../..").find_element(By.TAG_NAME, "input")
            el1.clear()
            time.sleep(0.5)
        el1 = el[0].find_element(By.XPATH, "../../..").find_element(By.TAG_NAME, "input")
        el1.send_keys("{0}".format(value))
        time.sleep(1)
        el2 = driver.find_elements(By.XPATH, ".//li[contains(concat(' ', @data-item-marker, ' '), '{0}')]".format(value))
        n = len(el2)
        for key in range(n):
            if el2[key].is_displayed():
                if el2[key].get_attribute('data-value') == "00000000-0000-0000-0000-000000000000":
                    text = 'Элемент \"' + value + '\"' + ' у поля \"' + name + '\" не найден'
                    write_file(text, "Error")
                    print(text)
                    screen_to_file(driver, 100)
                    raise MyException(text)
                else:
                    el2[key].click()
                    text = 'В поле \"' + name + '\"' + ' добавлено значение \"' + value + '\"'
                    write_file(text, "Good")
                    print(text)


#  Ввод поля с выбором колонки объекта
def enter_column_name(driver, name, column):

    el = driver.find_elements(By.XPATH, "//label[text() = '{0}' and @class = 't-label t-label-is-required ']".format(name))
    dl = len(el)
    if dl == 0:
        text = 'Поле \"' + name + '\" не найдено'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    if dl > 1:
        text = 'Не удалось однозначно определить поле \"' + name + '\"'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    else:
        el1 = el[0].find_element(By.XPATH, "../../..").find_element(By.TAG_NAME, "input")
        if el1.get_attribute('readonly'):
            text = 'Поле \"' + name + '\" недоступно для редактирования'
            write_file(text, "Error")
            print(text)
            screen_to_file(driver, 100)
            raise MyException(text)
        el[0].find_element(By.XPATH, "../../..").find_element(By.XPATH, ".//div[@class = 'base-edit-right-icon-wrapper']").click()
        el1 = driver.find_element(By.XPATH, ".//input[@placeholder = 'Выберите колонку']")
        el1.send_keys("{0}".format(column))
        time.sleep(2)
        el2 = driver.find_elements(By.XPATH, ".//li[contains(concat(' ', @data-item-marker, ' '), '{0}')]".format(column))
        n = len(el2)
        if n == 0:
            text = 'Элемент \"' + column + '\"' + ' у поля \"' + name + '\" не найден'
            write_file(text, "Error")
            print(text)
            screen_to_file(driver, 100)
            raise MyException(text)
        else:
            el2[0].click()
            click_button(driver, "Выбрать")
            text = 'В поле \"' + name + '\"' + ' добавлено значение \"' + column + '\"'
            write_file(text, "Good")
            print(text)



# Создание сервиса
def add_service(driver, service):
    el = driver.find_elements(By.XPATH, ".//div[@data-item-marker='{0}']".format(service))
    dl = len(el)
    if dl > 0:
        el[0].find_element(By.XPATH, ".//label[contains(text(), 'Получить')]").click()
        text = 'Выбран сервис \"' + service + '\"'
        write_file(text, "Good")
        print(text)
    else:
        text = 'Сервис \"' + service + '\" не найден'
        write_file(text, "Error")
        print('Сервис \"' + service + '\" не найден')
        screen_to_file(driver, 100)
        raise MyException(Exception)

# Разворачивает деталь, если скрыта
def open_detail(driver, detail):
    det = driver.find_elements(By.XPATH, "//div[@data-item-marker='{0}' and contains(concat(' ', @class, ' '), ' detail ')]".format(detail))
    dl = len(det)
    if dl == 0:
        text = 'Деталь \"' + detail + '\" не найдена'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    cl = det[0].get_attribute('class')
    dl = cl.find('ts-controlgroup-collapsed')
    if dl >= 0:
        ss = det[0].find_element(By.XPATH, "//span[contains(text(), '{0}')]".format(detail))
        driver.execute_script('arguments[0].scrollIntoView(true);', ss)
        ss.click()


# Нажимаем на добавить у детали
def click_add_detail(driver, detail):

    el = driver.find_elements(By.XPATH, ".//div[@data-item-marker='{0}' and contains("
                                       "concat(' ', @class, ' '), ' detail ')]".format(detail))
    dl = len(el)
    if dl == 0:
        text = 'Деталь \"' + detail + '\" не найдена'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    open_detail(driver, detail)

    el1 = driver.find_element(By.XPATH, ".//div[@data-item-marker= '{0}' and contains(concat(' ', @class, ' '), ' detail ')]".format(detail))
    el2 = el1.find_elements(By.XPATH, ".//span[contains(concat(' ', @data-item-marker, ' '), 'RecordButton')]")
    dl = len(el2)
    if dl == 0:
        text = 'Для детали \"' + detail + '\" невозможно добавить элемент'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    el2[0].click()
    text = 'Для детали \"' + detail + '\" Выбрали действие добавить'
    write_file(text, "Good")
    print(text)



# Нажимаем на тулбар у детали
def click_tool_detail(driver, detail):
    el = driver.find_elements(By.XPATH, ".//div[@data-item-marker='{0}' and contains(concat(' ', @class, ' '), ' detail ')]".format(detail))
    dl = len(el)
    if dl == 0:
        text = 'Деталь \"' + detail + '\" не найдена'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    open_detail(driver, detail)
    el1 = driver.find_element(By.XPATH, ".//div[@data-item-marker = '{0}' and contains(concat(' ', @class, ' '), ' detail ')]".format(detail))
    el2 = el1.find_elements(By.XPATH, ".//span[contains(concat(' ', @data-item-marker, ' '), 'ToolsButton')]")
    dl = len(el2)
    if dl == 0:
        text = 'Для детали \"' + detail + '\" недоступно меню \"Действия\"'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    el2[0].click()
    text = 'Для детали \"' + detail + '\" раскрыто меню \"Действия\"'
    write_file(text, "Good")
    print(text)


# Чтение данных в детали
def read_item_detail(driver, detail):
    el = driver.find_elements(By.XPATH, ".//div[@data-item-marker='{0}' and contains(concat(' ', @class, ' '), ' detail ')]".format(detail))
    dl = len(el)
    if dl == 0:
        text = 'Деталь \"' + detail + '\" не найдена'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    open_detail(driver, detail)
    el = driver.find_element(By.XPATH, ".//div[@data-item-marker='{0}' and contains(concat(' ', @class, ' '), ' detail ')]".format(detail))
    tag = 1
    while tag == 1:
        el1 = el.find_elements(By.XPATH, ".//span[contains(text(), 'Показать больше')]")
        dl = len(el1)
        if dl == 0:
           tag = 0
        else:
            if el1[0].is_displayed():
                el1[0].click()
                time.sleep(2)
            else:
                tag = 0
    el = driver.find_element(By.XPATH, ".//div[@data-item-marker='{0}' and contains(concat(' ', @class, ' '), ' detail ')]".format(detail))
    el1 = el.find_elements(By.XPATH, ".//div[contains(concat(' ', @tabindex, ' '), '0')]/div")
    #el2 = el1[0].find_elements(By.XPATH, "/div")
    line = len(el1)
    el2 = el1[0].find_elements(By.TAG_NAME, "div")
    column = len(el2)
    today = datetime.datetime.now()

    file = "./test_file/test_" + today.strftime("%d_%B") + '.docx'
    text = "Данные в детали \"" + detail + "\":"
    write_file(text, "Good")
    if os.path.exists(file):
        document = docx.Document(file)
    else:
        document = docx.Document()
    table = document.add_table(rows=line, cols=column)
    table.style = 'Table Grid'
    tt = el1[0].find_elements(By.TAG_NAME, "div")
    for col in range(column):
        value = tt[col].get_attribute('data-item-marker')
        # получаем ячейку таблицы
        cell = table.cell(0, col)
        # записываем в ячейку данные
        cell.text = value
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
    for row in range(line-1):
        tt = el1[row + 1].find_elements(By.XPATH, "*")
        for col in range(column):
            tt1 = tt[col].find_elements(By.TAG_NAME, "span")
            dl = len(tt1)
            if dl > 0:
                value = tt1[0].text
                cell = table.cell(row + 1, col)
                cell.text = value
    document.save(file)






# Ищем деталь
def search_detail(driver, detail):
    el = driver.find_elements(By.XPATH, ".//div[@data-item-marker='{0}' and contains(concat(' ', @class, ' '), ' detail ')]".format(detail))
    dl = len(el)
    if dl == 0:
        text = 'Деталь \"' + detail + '\" не найдена'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    el = driver.find_element(By.XPATH, ".//div[@data-item-marker='{0}' and contains(concat(' ', @class, ' '), ' detail ')]".format(detail))
    return el



# Выбрать первый элемент в детали
def select_first_object_detail(driver, detail):
    el = driver.find_elements(By.XPATH, ".//div[@data-item-marker='{0}' and contains(concat(' ', @class, ' '), ' detail ')]".format(detail))
    dl = len(el)
    if dl == 0:
        text = 'Деталь \"' + detail + '\" не найдена'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    el = driver.find_element(By.XPATH, ".//div[@data-item-marker='{0}' and contains(concat(' ', @class, ' '), ' detail ')]".format(detail))
    el2 = el.find_element(By.XPATH, ".//div[contains(concat(' ', @id, ' '), 'DetailV2DataGridGrid-')]")
    el3 = el2.find_elements(By.XPATH, ".//div[contains(concat(' ', @id, ' '), 'DetailV2DataGridGrid-')]")
    dl = len(el3)
    if dl == 0:
        text = 'В детали \"' + detail + '\" нет объектов'
        write_file(text, "Error")
        print(text)
        #raise MyException(Exception)
    else:
        el2.click()
        #click_tool_detail(driver, 'Карьера')
        #action_tools_button(driver, 'Удалить')
        #message_box(driver, "Да")
        text = 'В детали \"' + detail + '\" выбран первый элемент'
        write_file(text, "Good")
        print(text)



# Нажать кнопку
def click_button(driver, button):
    el = driver.find_elements(By.XPATH, ".//span[contains(text(), '{0}')]".format(button))
    dl = len(el)
    if dl == 0:
        text = 'Кнопка \"' + button + '\" не найдена'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    if el[0].is_displayed():
        driver.find_element(By.XPATH, ".//span[contains(text(), '{0}')]".format(button)).click()
        text = 'Кнопка \"' + button + '\" нажата'
        write_file(text, "Good")
        print(text)
    else:
        text = 'Кнопка \"' + button + '\" не кликабельна'
        write_file(text, 2)
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    time.sleep(0.5)


# Выбор элемента в окне
def select_window(driver, name, value):
    el = driver.find_element(By.ID, "columnEdit-el")
    el.clear()
    el.send_keys("{0}".format(name))
    time.sleep(1)
    el2 = driver.find_elements(By.XPATH, ".//li[contains(concat(' ', @data-item-marker, ' '), '{0}')]".format(name))
    n = len(el2)
    for key in range(n):
        if el2[key].is_displayed():
            if el2[key].get_attribute('data-value') == "00000000-0000-0000-0000-000000000000":
                text = 'Элемент \"' + name + '\" не найден'
                write_file(text, "Error")
                print('Элемент \"' + name + '\" не найден')
                screen_to_file(driver, 100)
                raise MyException(text)
            else:
                el2[key].click()
                break
    el = driver.find_element(By.XPATH, ".//div[@data-item-marker='searchEdit']/input[1]")
    el.clear()
    el.send_keys("{0}".format(value))
    click_button(driver, "Поиск")
    el = driver.find_elements(By.XPATH, ".//span[contains(text(), '{0}')]".format(value))
    dl = len(el)
    if dl == 0:
        text = 'Элемент \"' + value + '\" не найден'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
       # "//table[@id='abc']//div/nobr[.='abc']/../.."
    el = driver.find_element(By.ID, "grid-grid-wrap").find_element(By.XPATH,
        ".//span[contains(text(), '{0}')]".format(value))
    el1 = el.find_element(By.XPATH, "..")
    el2 = el1.find_element(By.XPATH, "..")
    el2.click()
    text = 'Добавлен \"' + value + '\"'
    write_file(text, "Good")
    print(text)
    time.sleep(0.5)


# Открытие вкладки
def open_tab(driver, tab):
    driver.find_element(By.TAG_NAME, 'body').send_keys(Keys.HOME)
    time.sleep(2)
    page_lock(driver)
    el = driver.find_elements(By.XPATH, "//ul[@class='ts-tabpanel-items scroll-animation ']/li[text() = '{0}']".format(tab))
    dl = len(el)
    if dl == 0:
        text = 'Вкладка \"' + tab + '\" не найдена'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    if dl > 1:
        text = 'Найдено несколько подходящих вкладок \"' + tab + '\"'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    if not el[0].is_displayed():
        loctab = el[0].location
        scr = driver.find_element(By.ID, "CasePageTabsTabPanel-scroll-right")
        scr_left = scr.location
        if int(loctab["x"]) > (scr_left["x"]):
            while not el[0].is_displayed():
                driver.find_element(By.CLASS_NAME, "ts-tabpanel-scroll-right").click()
                time.sleep(0.5)
                el = driver.find_elements(By.XPATH, "//ul[@class='ts-tabpanel-items scroll-animation ']/li[text() = '{0}']".format(tab))
        else:
            while not el[0].is_displayed():
                driver.find_element(By.CLASS_NAME, "ts-tabpanel-scroll-left").click()
                time.sleep(0.5)
                el = driver.find_elements(By.XPATH, "//ul[@class='ts-tabpanel-items scroll-animation ']/li[text() = '{0}']".format(tab))
    text = 'Переходим на вкладку \"' + tab + '\"'
    write_file(text, "Good")
    print(text)
    el[0].click()


def visa_approve(driver, define):
    wait = WebDriverWait(driver, 10)
    # el = driver.find_element(By.XPATH, ".//div[contains(concat(' ', @id, ' '), 'VisaDetailV2DataGridGrid-wrap')]")
    el1 = driver.find_elements(By.XPATH, "//span[contains(text(), 'Ожидает визирования')]")
    dl = len(el1)
    if bool(dl):
        time.sleep(2)
        el = wait.until(ec.element_to_be_clickable((By.XPATH, "//span[contains(text(), 'Ожидает визирования')]")))
        # el1 = el.find_element(By.XPATH, "//span[contains(text(), 'Ожидает визирования')]")
        driver.execute_script('arguments[0].scrollIntoView(true);', el)
        el.click()
        driver.find_element(By.XPATH,
            "//div[contains(concat(' ', @id, ' '), 'VisaDetailV2DetailControlGroup-tools')]").find_element(By.XPATH,
            ".//span[contains(concat(' ', @data-item-marker, ' '), 'ToolsButton')]").click()
        driver.find_element(By.XPATH,
            ".//li[contains(concat(' ', @data-item-marker, ' '), '{0}')]".format(define)).click()
        print("Для визы выбрано действие:" + define)
    else:
        print("Нет согласований, ожидающих визирования пользователя")

# Смотрим, нет ли перекрытия страницы, если есть, то ждём
def page_lock(driver):
    WebDriverWait(driver, 60).until(ec.invisibility_of_element_located(
        (By.XPATH, ".//div[contains(concat(' ', @class, ' '), 'ts-mask-spinner-caption')]")))



# Выбор действия из списка
def action_tools_button(driver, action):
    time.sleep(1)
    el1 = driver.find_elements(By.XPATH, ".//li[@data-item-marker = '{0}']".format(action))
    if len(el1) == 0:
        text = 'Действие \"' + action + '\" недоступно'
        write_file(text, "Error")
        print('Действие \"' + action + '\" недоступно')
        screen_to_file(driver, 100)
        raise MyException(text)
    else:
        driver.find_element(By.XPATH, ".//li[@data-item-marker = '{0}']".format(action)).click()
        text = 'Действие \"' + action + '\" активировано'
        write_file(text, "Good")
        print('Действие \"' + action + '\" активировано')


# Открываем реестр
def open_object(driver, name):
    cleaning_filters(driver)
    create_filters(driver, name)
    wait = WebDriverWait(driver, 10)
    el = driver.find_elements(By.XPATH, "//div[@data-item-marker = '{0}']".format(name))
    if len(el) == 0:
        text = '\"' + name + '\" не найден'
        write_file(text, "Error")
        print(text)
        screen_to_file(driver, 100)
        raise MyException(text)
    el1 = wait.until(ec.element_to_be_clickable((By.XPATH, ".//a[contains(concat(' ', @title, ' '), '{0}')]".format(name))))
    url = el1.get_attribute('href')
    text = 'Открыт объект \"' + name + '\"'
    write_file(text, "Good")
    print(text)
    driver.get(url)

# Открыть дашборды
def open_analitics(driver):
    el = driver.find_element(By.ID, "view-button-AnalyticsDataView-wrapperEl")
    el.click()
# открыть список объектов в разделе
def open_register(driver):
    el = driver.find_element(By.ID, "view-button-GridDataView-wrapperEl")
    el.click()

# Всплывающее окошко
def message_box(driver, yes_no):
    time.sleep(2)
    el = driver.find_elements(By.XPATH, "//div[@class = 'ts-messagebox-caption']")
    dl = len(el)
    if dl > 0:
        text = el[0].text
        text = "Появилось окошко с сообщением:\n" + text
        write_file(text, "Good")
        screen_to_file(driver, 50)
        print(text)
        el1 = el[0].find_elements(By.XPATH, "//span[@data-item-marker = '{0}']".format(yes_no))
        dl = len(el1)
        if dl > 0:
            el[0].find_element(By.XPATH, "//span[@data-item-marker = '{0}']".format(yes_no)).click()
            text = "Выбрано: " + yes_no
            write_file(text, "Good")
            print(text)
        page_lock(driver)

#  Скриншот центар экрана (100 - весь экран, 0 - точка в центре экрана)
def screenshot(driver, size):
    if size < 0:
        size = 0
    elif size >100:
        size = 100
    Height = str(driver.execute_script('return document.documentElement.clientHeight'))
    Width = str(driver.execute_script('return document.documentElement.clientWidth'))
    Height = int(Height)
    Width = int(Width)
    x1 = float(Width)/2*(1-float(size)/100)
    y1 = float(Height)/2*(1-float(size)/100)
    x2 = float(Width)/2*(1+float(size)/100)
    y2 = float(Height)/2*(1+float(size)/100)
    img = ImageGrab.grab((int(x1), int(y1), int(x2), int(y2)))
    today = datetime.datetime.now()
    texttoday_data = today.strftime("%B-%m-%d-%H.%M.%S")
    namescreen = "../screen/screen_" + texttoday_data + ".bmp"
    if not os.path.exists("../screen/"):
        os.mkdir("../screen/")
    img.save(namescreen, "BMP")
    return namescreen

def write_file(text,err):
    today = datetime.datetime.now()
    print(os.getcwd())
    file = "../test_log/test_" + today.strftime("%d_%B") + '.docx'

    if err == "Error":
        if os.path.exists(file):
            document = docx.Document(file)
            par = document.add_paragraph()
            par.add_run(text).font.color.rgb = RGBColor(0xff, 0, 0)
            document.save(file)
        else:
            if not os.path.exists("../test_log/"):
                os.mkdir("../test_log/")
            document = docx.Document()
            par = document.add_paragraph()
            par.add_run(text).font.color.rgb = RGBColor(0xff, 0, 0)
            document.save(file)
    if err == "Good":
        if os.path.exists(file):
            document = docx.Document(file)
            par = document.add_paragraph()
            par.add_run(text)
            document.save(file)
        else:
            if not os.path.exists("../test_log/"):
                os.mkdir("../test_log/")
            document = docx.Document()
            par = document.add_paragraph()
            par.add_run(text)
            document.save(file)
    if err == 2:
        if os.path.exists(file):
            document = docx.Document(file)
            par = document.add_paragraph()
            par.add_run(text).font.color.rgb = RGBColor(0, 0, 0xff)
            document.save(file)
        else:
            if not os.path.exists("../test_log/"):
                os.mkdir("../test_log/")
            document = docx.Document()
            par = document.add_paragraph()
            par.add_run(text).font.color.rgb = RGBColor(0, 0, 0xff)
            document.save(file)
    if err == "Test":
        if os.path.exists(file):
            document = docx.Document(file)
            par = document.add_paragraph()
            par.add_run(text).font.color.rgb = RGBColor(0, 0xff, 0)
            par.alignment = 1
            par.bold = 1
            document.save(file)
        else:
            if not os.path.exists("../test_log/"):
                os.mkdir("../test_log/")
            document = docx.Document()
            par = document.add_paragraph()
            par.add_run(text).font.color.rgb = RGBColor(0,  0xff, 0)
            par.bold = 1
            par.alignment = 1
            document.save(file)


def screen_to_file(driver, size):
    img = screenshot(driver, size)
    today = datetime.datetime.now()
    file = "../test_log/test_" + today.strftime("%d_%B") + '.docx'
    if os.path.exists(file):
        document = docx.Document(file)
        document.add_picture(img, width=Inches(7.0))
        document.save(file)
    else:
        if not os.path.exists("../test_log/"):
            os.mkdir("../test_log/")
        document = docx.Document()
        document.add_picture(img, width=Inches(7.0))
        document.save(file)


def to_latinica(text):
    return cyrtranslit.to_latin(text)





# Активация действия
def action_selection(driver, action):
    time.sleep(3)
    driver.find_element(By.XPATH, "//span[@data-item-marker = 'actions']").click()
    # driver.find_element(By.XPATH, ".//span[contains(text(), 'Действия')]").click()
    el = driver.find_elements(By.XPATH, ".//li[@data-item-marker = '{0}']".format(action))
    if len(el) == 0:
        text = 'Действие \"' + action + '\" недоступно'
        write_file(text, "Error")
        print('Действие \"' + action + '\" недоступно')
        # raise MyException(Exception)
    else:
        driver.find_element(By.XPATH, ".//li[@data-item-marker = '{0}']".format(action)).click()
        text = 'Действие \"' + action + '\" активировано'
        write_file(text, "Good")
        print('Действие \"' + action + '\" активировано')

# Скрин последнего письма на почте mail
def login_mail(driver, name, passw):

    driver.get('https://mail.ru/')
    wait = WebDriverWait(driver, 30)
    driver.maximize_window()
    user_name = driver.find_element(By.ID, "mailbox:login")
    user_name.clear()
    user_name.send_keys(name)
    driver.find_element(By.ID, "mailbox:submit").click()
    password = driver.find_element(By.ID, "mailbox:password")
    password.clear()
    password.send_keys(passw)
    driver.find_element(By.ID, "mailbox:submit").click()
    el = wait.until(ec.element_to_be_clickable((By.XPATH, ".//div[contains(text(), 'Входящие')]")))
    el.click()
    el = wait.until(ec.element_to_be_clickable((By.XPATH, "//div[@class='dataset__items']/a[1]")))
    el.click()
    wait.until(ec.invisibility_of_element_located((By.XPATH, ".//span[contains(text(), 'Переслать')]")))
    driver.save_screenshot("screenshotmal.png")

# Поиск в новом фрейме и возврат обратно
    #iframe = driver.find_elements(By.TAG_NAME, 'iframe')[0]
    #driver.switch_to_frame(iframe)
    #driver.find_element(By.XPATH, "//div").send_keys("Текст для проверки заполненя поля новости")
    #driver.switch_to.default_content()



    #ss = driver.find_element(By.XPATH, ".//ul[contains(concat(' ', @id, ' '), 'PageTabsTabPanel-tabpanel-items')]").find_element(By.XPATH, ".//li[contains(text(), '{0}')]".format(tab))
    #driver.execute_script('arguments[0].scrollIntoView(true);', ss)
    #wait.until(ec.element_to_be_clickable((By.XPATH, ".//li[contains(text(), '{0}')]".format(tab))))
    #ActionChains(driver).move_to_element(
     #   driver.find_element(By.XPATH, ".//li[contains(text(), '{0}')]".format(tab))).perform()

    #driver.find_element(By.XPATH, ".//ul[contains(concat(' ', @id, ' '), 'PageTabsTabPanel-tabpanel-items')]").find_element(By.XPATH,
    #    ".//li[contains(text(), '{0}')]".format(tab)).click()
    # Дяля раздела Контакты
    #driver.find_element(By.XPATH, ".//ul[contains(concat(' ', @id, ' '), 'ContactPageV2TabsTabPanel-tabpanel-items')]").find_element(By.XPATH, ".//li[contains(text(), '{0}')]".format(tab)).click()
    # Для остальных разделов
    #driver.find_element(By.XPATH, ".//div[contains(concat(' ', @id, ' '), 'PageTabsTabPanel-items-wrap')]").find_element(By.XPATH, ".//li[contains(text(), '{0}')]".format(tab)).click()
    #text = 'Переходим на вкладку \"' + tab + '\"'
    #write_file(text, "Good")
    #print('Переходим на вкладку \"' + tab + '\"')



