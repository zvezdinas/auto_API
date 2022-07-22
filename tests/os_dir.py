import time
import datetime
import os
import docx
from docx.dml.color import ColorFormat
from docx.shared import RGBColor
from docx.shared import Inches
from docx.text.run import Font, Run
from PIL import ImageGrab
import cyrtranslit
from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as ec
from selenium.webdriver.support.ui import WebDriverWait

today = datetime.datetime.now()
print(os.getcwd())
file = "../screen/test_" + today.strftime("%d_%B") + '.docx'

text = 'text text text text text text text text text text '

if os.path.exists(file):
    document = docx.Document(file)
    par = document.add_paragraph()
    par.add_run(text).font.color.rgb = RGBColor(0xff, 0, 0)
    document.save(file)
else:
    if not os.path.exists("../screen/"):
        os.mkdir("../screen/")

    document = docx.Document()
    par = document.add_paragraph()
    par.add_run(text).font.color.rgb = RGBColor(0xff, 0, 0)
    document.save(file)